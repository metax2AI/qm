#!/usr/bin/env python3
"""Regenerate the .xlsx, .docx and .pdf demo files from their .csv/.md sources.

Run this after editing any source file — most often after moving the follow-up
dates to the demo day, which the demo-data README asks for. Regenerating is what
keeps the formats from drifting apart, since a contradiction between them is the
signal we rely on to catch a parser reading something wrong.

    pip install openpyxl python-docx reportlab fonttools
    python deploy/layers/metax2ai/scripts/generate-demo-office-formats.py

The PDF embeds Noto Sans SC (SIL Open Font License) so the Chinese text renders
on readers without an Adobe CJK font pack. The font is downloaded to a cache
directory rather than committed; ReportLab subsets it, so the PDF stays small.
"""

import csv
import datetime as dt
import re
import sys
import urllib.request
from pathlib import Path

from docx import Document
from docx.shared import Pt
from fontTools import ttLib
from fontTools.varLib import instancer
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font
from openpyxl.utils import get_column_letter
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

DEMO = Path(__file__).resolve().parent.parent / "demo-data"
CACHE = Path.home() / ".cache" / "qm-demo-fonts"
FONT_URL = "https://github.com/google/fonts/raw/main/ofl/notosanssc/NotoSansSC%5Bwght%5D.ttf"
FONT_NAME = "NotoSansSC"

DATE_COLUMNS = {"下次跟进日期", "最近跟进日期"}
NUMBER_COLUMNS = {"预计金额(元)"}


def cjk_font() -> str:
    CACHE.mkdir(parents=True, exist_ok=True)
    variable = CACHE / "NotoSansSC-Variable.ttf"
    static = CACHE / "NotoSansSC-Regular.ttf"

    if not static.exists():
        if not variable.exists():
            last = None
            for attempt in range(1, 6):
                try:
                    print(f"downloading {FONT_URL} (attempt {attempt})", file=sys.stderr)
                    with urllib.request.urlopen(FONT_URL, timeout=120) as response:
                        payload = response.read()
                    declared = response.headers.get("Content-Length")
                    if declared and len(payload) != int(declared):
                        raise OSError(f"truncated: got {len(payload)} of {declared} bytes")
                    variable.write_bytes(payload)
                    break
                except OSError as error:
                    last = error
            else:
                raise SystemExit(f"could not download the CJK font: {last}")

        # The variable font defaults to wght=100 (Thin), which is too faint for a
        # document a customer reads. Pin the regular weight before embedding.
        font = ttLib.TTFont(str(variable))
        instancer.instantiateVariableFont(font, {"wght": 400}).save(str(static))

    pdfmetrics.registerFont(TTFont(FONT_NAME, str(static)))
    return FONT_NAME


def build_xlsx() -> None:
    with (DEMO / "客户跟进表.csv").open(encoding="utf-8", newline="") as fh:
        rows = list(csv.reader(fh))
    header, body = rows[0], rows[1:]

    wb = Workbook()
    ws = wb.active
    ws.title = "客户跟进"
    ws.append(header)
    for cell in ws[1]:
        cell.font = Font(bold=True)
        cell.alignment = Alignment(vertical="center")
    ws.freeze_panes = "A2"

    for record in body:
        typed = []
        for name, value in zip(header, record):
            if not value:
                typed.append(None)
            elif name in DATE_COLUMNS:
                typed.append(dt.date.fromisoformat(value))
            elif name in NUMBER_COLUMNS:
                typed.append(int(value))
            else:
                typed.append(value)
        ws.append(typed)

    for index, name in enumerate(header, start=1):
        letter = get_column_letter(index)
        if name in DATE_COLUMNS:
            for cell in ws[letter][1:]:
                cell.number_format = "yyyy-mm-dd"
        if name in NUMBER_COLUMNS:
            for cell in ws[letter][1:]:
                cell.number_format = "#,##0"
        widest = max([len(name)] + [len(r[index - 1]) for r in body])
        ws.column_dimensions[letter].width = min(46, max(12, widest * 1.6))

    wb.save(DEMO / "客户跟进表.xlsx")


def build_docx() -> None:
    src = DEMO / "会议纪要" / "2026-08-06-重点客户跟进会.md"
    doc = Document()
    doc.styles["Normal"].font.name = "Microsoft YaHei"
    doc.styles["Normal"].font.size = Pt(10.5)

    for line in src.read_text(encoding="utf-8").splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("- "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        else:
            doc.add_paragraph(stripped)

    doc.save(DEMO / "会议纪要" / "2026-08-06-重点客户跟进会.docx")


def md_table_rows(block: list[str]) -> list[list[str]]:
    rows = []
    for line in block:
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        if all(re.fullmatch(r":?-+:?", c) for c in cells):
            continue
        rows.append(cells)
    return rows


def build_pdf() -> None:
    font = cjk_font()
    src = DEMO / "产品资料" / "星云客户云平台产品手册.md"

    base = getSampleStyleSheet()
    body = ParagraphStyle(
        "CjkBody",
        parent=base["BodyText"],
        fontName=font,
        fontSize=10.5,
        leading=17,
        alignment=TA_LEFT,
        wordWrap="CJK",
    )
    h1 = ParagraphStyle("CjkH1", parent=base["Heading1"], fontName=font, fontSize=18, leading=26, spaceAfter=10)
    h2 = ParagraphStyle(
        "CjkH2", parent=base["Heading2"], fontName=font, fontSize=13, leading=20, spaceBefore=12, spaceAfter=6
    )
    h3 = ParagraphStyle(
        "CjkH3", parent=base["Heading3"], fontName=font, fontSize=11.5, leading=18, spaceBefore=9, spaceAfter=4
    )

    flow = []
    pending_table: list[str] = []

    def flush_table() -> None:
        if not pending_table:
            return
        rows = md_table_rows(pending_table)
        width = max(len(r) for r in rows)
        usable = 166 * mm
        col_widths = [42 * mm] + [(usable - 42 * mm) / (width - 1)] * (width - 1) if width > 1 else [usable]
        data = [[Paragraph(c, body) for c in row] for row in rows]
        table = Table(data, colWidths=col_widths, repeatRows=1)
        table.setStyle(
            TableStyle(
                [
                    ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#9aa0a6")),
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#eef1f5")),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 5),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                    ("TOPPADDING", (0, 0), (-1, -1), 4),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                ]
            )
        )
        flow.append(table)
        flow.append(Spacer(1, 8))
        pending_table.clear()

    for line in src.read_text(encoding="utf-8").splitlines():
        stripped = line.strip()
        if stripped.startswith("|"):
            pending_table.append(stripped)
            continue
        flush_table()
        if not stripped:
            continue
        if stripped.startswith("### "):
            flow.append(Paragraph(stripped[4:], h3))
        elif stripped.startswith("## "):
            flow.append(Paragraph(stripped[3:], h2))
        elif stripped.startswith("# "):
            flow.append(Paragraph(stripped[2:], h1))
        elif stripped.startswith("- "):
            flow.append(Paragraph("• " + stripped[2:], body))
        else:
            flow.append(Paragraph(stripped, body))
    flush_table()

    SimpleDocTemplate(
        str(DEMO / "产品资料" / "星云客户云平台产品手册.pdf"),
        pagesize=A4,
        title="星云客户云平台产品手册",
        author="星云云",
        leftMargin=22 * mm,
        rightMargin=22 * mm,
        topMargin=20 * mm,
        bottomMargin=20 * mm,
    ).build(flow)


build_xlsx()
build_docx()
build_pdf()
print("regenerated 客户跟进表.xlsx, 2026-08-06-重点客户跟进会.docx, 星云客户云平台产品手册.pdf")
